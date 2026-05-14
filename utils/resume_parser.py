"""
Resume Parser Utility
Extracts raw text from uploaded PDF and DOCX files
Lightweight: uses pdfplumber (better than PyPDF2) and python-docx
"""

import os
import re


def extract_text_from_file(filepath: str) -> str:
    """
    Extract plain text from a resume file.
    Supports PDF and DOCX formats.
    Returns empty string on failure.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        return _extract_from_pdf(filepath)
    elif ext in ('.docx', '.doc'):
        return _extract_from_docx(filepath)
    else:
        return ''


def _extract_from_pdf(filepath: str) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return '\n'.join(text_parts)
    except ImportError:
        # Fallback to PyPDF2 if pdfplumber not available
        try:
            import PyPDF2
            text_parts = []
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or '')
            return '\n'.join(text_parts)
        except Exception as e:
            return f'[Error extracting PDF: {str(e)}]'
    except Exception as e:
        return f'[Error extracting PDF: {str(e)}]'


def _extract_from_docx(filepath: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return '\n'.join(paragraphs)
    except Exception as e:
        return f'[Error extracting DOCX: {str(e)}]'


def clean_text(text: str) -> str:
    """Normalize extracted text: remove excess whitespace, fix encoding."""
    text = re.sub(r'\s+', ' ', text)           # Collapse whitespace
    text = re.sub(r'[^\x00-\x7F]+', ' ', text) # Remove non-ASCII
    return text.strip()


def count_words(text: str) -> int:
    """Count words in resume text."""
    return len(text.split())


def extract_email(text: str) -> str:
    """Extract email address from resume text."""
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    match = re.search(pattern, text)
    return match.group(0) if match else ''


def extract_phone(text: str) -> str:
    """Extract phone number from resume text."""
    pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    match = re.search(pattern, text)
    return match.group(0) if match else ''


def extract_sections(text: str) -> dict:
    """
    Attempt to identify resume sections by common headings.
    Returns a dict of section_name -> content.
    """
    section_keywords = {
        'education': ['education', 'academic', 'qualification', 'degree'],
        'experience': ['experience', 'work experience', 'internship', 'employment', 'projects', 'project'],
        'skills': ['skills', 'technical skills', 'technologies', 'tools', 'competencies'],
        'achievements': ['achievements', 'awards', 'honors', 'accomplishments', 'certifications', 'certificates'],
        'summary': ['summary', 'objective', 'profile', 'about me', 'overview'],
    }

    lines = text.split('\n')
    sections = {}
    current_section = 'other'
    sections[current_section] = []

    for line in lines:
        line_lower = line.lower().strip()
        matched = False
        for section, keywords in section_keywords.items():
            if any(kw in line_lower for kw in keywords) and len(line.strip()) < 50:
                current_section = section
                sections.setdefault(current_section, [])
                matched = True
                break
        if not matched and line.strip():
            sections.setdefault(current_section, []).append(line.strip())

    return {k: ' '.join(v) for k, v in sections.items()}
